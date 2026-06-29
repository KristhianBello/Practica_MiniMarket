from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Cardenas_Mejia_Intriago_Bello_Tema3_2_APIPropia.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C2CC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", size=11, bold=False, color=None, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 8)
    pf.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 12.5, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"- {item}")
        set_run_font(r, size=10.5)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=font_size, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    return table


def add_placeholder(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_table_borders(table, color="AAB7C4", size="6")
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(caption)
    set_run_font(r, size=9.5, bold=True, color=MUTED)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "Aplicaciones Moviles Hibridas | Tema 3.2"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("MiniMarket App - API propia con JSON Server")
    set_run_font(run, size=9, color=MUTED)
    return doc


def build_doc():
    doc = setup_document()

    # Cover / masthead
    add_para(doc, "UNIVERSIDAD LAICA ELOY ALFARO DE MANABI", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "FACULTAD CIENCIAS DE LA VIDA Y TECNOLOGIAS", size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "CARRERA DE SOFTWARE", size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, "TRABAJO ACADEMICO ASINCRONICO", size=15, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "Tema 3.2: Consumo de Servicios Web Propios", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

    metadata = [
        ("Asignatura", "Aplicaciones Moviles Hibridas"),
        ("Unidad", "Unidad 3. Gestion de Estado e Interaccion con Datos"),
        ("Nivel / paralelo", "Septimo semestre"),
        ("Docente", "Ing. Cesar Stalin Villavicencio Palacios, Mg."),
        ("Proyecto", "MiniMarket App con API REST propia"),
        ("Fecha", "29 de junio de 2026"),
    ]
    add_table(doc, ["Campo", "Detalle"], metadata, [1900, 7460], header_fill=LIGHT_BLUE, font_size=9.5)
    add_para(doc, "Integrantes", size=11, bold=True, color=DARK_BLUE, before=8, after=3)
    add_bullets(
        doc,
        [
            "Cardenas Avila Emilio Sleimen",
            "Anthony Axel Mejia Ordonez",
            "Michael Agustin Intriago Benitez",
            "Kristhian Augusto Bello Soledispa",
        ],
    )

    add_heading(doc, "1. Descripcion del proyecto y objetivo", 1)
    add_para(
        doc,
        "El proyecto MiniMarket App fue adaptado como una aplicacion movil hibrida desarrollada con Expo, React Native y Expo Router. Su objetivo es consumir una API REST propia creada con JSON Server, mostrar productos en pantalla, manejar estados de carga/error, permitir actualizacion de datos y registrar nuevos productos mediante una peticion POST.",
        size=10.5,
    )
    add_para(
        doc,
        "La aplicacion conserva el enfoque de mini tienda: catalogo de productos, detalle, carrito y resumen de compra. La diferencia principal frente a la version inicial es que el catalogo ya no depende de datos quemados en el codigo, sino del endpoint REST propio definido en la carpeta api-productos.",
        size=10.5,
    )

    add_heading(doc, "2. API propia creada con JSON Server", 1)
    add_para(
        doc,
        "La API se encuentra en api-productos/db.json y expone el recurso principal productos. El servidor se ejecuta con npm run api, usando json-server db.json --host 0.0.0.0 --port 3001 para permitir consumo desde Expo Go en un celular conectado a la misma red Wi-Fi.",
        size=10.5,
    )
    api_rows = [
        ("Recurso", "productos"),
        ("Endpoint PC", "http://localhost:3001/productos"),
        ("Endpoint Expo Go", "http://192.168.10.110:3001/productos"),
        ("Campos", "id, nombre, precio, categoria, stock, descripcion, detalle, integrante, imageKey"),
        ("Finalidad", "Administrar articulos tecnologicos del MiniMarket y evidenciar GET/POST."),
    ]
    add_table(doc, ["Elemento", "Detalle del proyecto"], api_rows, [1850, 7510], header_fill=LIGHT_BLUE, font_size=9)

    add_heading(doc, "3. Productos personalizados", 1)
    product_rows = [
        ("1", "Audifonos Bluetooth JBL", "Audio", "$25.00", "10"),
        ("2", "Mouse inalambrico ergonomico", "Perifericos", "$15.00", "15"),
        ("3", "Teclado mecanico compacto", "Perifericos", "$45.00", "8"),
        ("4", "Cargador rapido USB-C 65W", "Accesorios", "$12.00", "20"),
        ("5", "Soporte ajustable para celular", "Accesorios", "$18.00", "12"),
        ("6", "Cable USB-C reforzado", "Accesorios", "$7.50", "18"),
        ("7", "Pad mouse gamer antideslizante", "Perifericos", "$9.99", "14"),
        ("8", "Mini parlante Bluetooth", "Audio", "$22.50", "9"),
    ]
    add_table(doc, ["ID", "Producto", "Categoria", "Precio", "Stock"], product_rows, [500, 4060, 1800, 1400, 1200], header_fill=LIGHT_GRAY, font_size=8.6)
    add_para(doc, "Ademas, durante las pruebas se registro un producto nuevo desde la app, demostrando la peticion POST hacia /productos.", size=9.5, italic=True, color=MUTED, after=2)

    add_heading(doc, "4. Endpoints consumidos", 1)
    endpoint_rows = [
        ("GET", "/productos", "Lista todos los productos.", "Catalogo principal y captura en navegador."),
        ("GET", "/productos/:id", "Consulta un producto especifico.", "Pantalla Detalle del producto."),
        ("POST", "/productos", "Crea un producto nuevo.", "Formulario Registrar producto en el catalogo."),
    ]
    add_table(doc, ["Metodo", "Endpoint", "Descripcion", "Evidencia"], endpoint_rows, [900, 1900, 3300, 3260], header_fill=LIGHT_BLUE, font_size=8.8)

    add_heading(doc, "5. Consumo desde la app movil", 1)
    add_para(
        doc,
        "La capa de consumo HTTP esta centralizada en src/api/products.ts. La funcion fetchProducts() realiza GET /productos; fetchProductById(id) consulta GET /productos/:id; y createProduct(input) envia POST /productos con cabecera Content-Type: application/json y cuerpo JSON.stringify(). Antes de procesar la respuesta, el codigo valida response.ok para detectar errores HTTP.",
        size=10.2,
    )
    add_para(
        doc,
        "La pantalla src/app/(tabs)/index.tsx carga los productos con FlatList, muestra ActivityIndicator durante la carga, presenta un mensaje de error si no se puede conectar, permite pull-to-refresh con RefreshControl y contiene un formulario con etiquetas para nombre, precio, stock, categoria e integrante que registra.",
        size=10.2,
    )

    add_heading(doc, "6. Evidencias solicitadas", 1)
    add_para(doc, "Los siguientes espacios estan preparados para pegar las capturas reales antes de entregar el archivo final en Moodle.", size=9.8, italic=True, color=MUTED)
    evidence_rows = [
        ("1", "Servidor JSON Server ejecutandose", "Terminal con npm run api y puerto 3001."),
        ("2", "Endpoint en navegador", "http://192.168.10.110:3001/productos abierto desde el celular o PC."),
        ("3", "Catalogo en la app", "MiniMarket mostrando productos consumidos desde la API."),
        ("4", "Registro POST", "Formulario con campos llenos y producto nuevo en la lista."),
        ("5", "Detalle de producto", "Pantalla que consulta GET /productos/:id."),
    ]
    add_table(doc, ["No.", "Captura requerida", "Que debe mostrar"], evidence_rows, [650, 3200, 5510], header_fill=LIGHT_GRAY, font_size=8.8)
    add_placeholder(doc, "PEGAR CAPTURA 1: JSON Server ejecutandose")
    add_placeholder(doc, "PEGAR CAPTURA 2: Endpoint /productos en navegador")
    add_placeholder(doc, "PEGAR CAPTURA 3: App MiniMarket mostrando catalogo y/o POST")

    add_heading(doc, "7. Reflexion tecnica final", 1)
    reflection = (
        "La principal dificultad fue conectar Expo Go con la API propia cuando la app se ejecuta en un celular fisico. "
        "Al inicio, la app no podia usar localhost porque en el telefono esa direccion apunta al propio dispositivo, no a la computadora. "
        "Tambien aparecio una IP incorrecta 169.254.x.x, por lo que se forzo la IP real de la red local 192.168.10.110. "
        "La solucion fue levantar JSON Server con --host 0.0.0.0, probar el endpoint desde el navegador del celular y configurar EXPO_PUBLIC_API_URL. "
        "Con esto la app pudo consumir GET /productos y registrar nuevos datos con POST /productos. "
        "Aprendimos que una API REST requiere endpoint, metodo HTTP, cabeceras, cuerpo y validacion de respuesta. "
        "Tambien fue importante manejar ActivityIndicator, mensajes de error y RefreshControl para mejorar la experiencia del usuario. "
        "El proyecto evidencia integracion real entre backend local y app movil hibrida."
    )
    add_para(doc, reflection, size=10.3)

    add_heading(doc, "8. Preguntas de reflexion", 1)
    qa_rows = [
        ("Por que no siempre funciona localhost?", "Porque en un celular fisico localhost apunta al propio telefono. Debe usarse la IP de la computadora donde corre JSON Server."),
        ("Diferencia entre GET y POST", "GET consulta datos existentes; POST envia un cuerpo JSON para crear un nuevo registro."),
        ("Por que validar response.ok?", "Evita intentar procesar como JSON una respuesta de error y permite mostrar un mensaje claro al usuario."),
        ("Que hacer ante Network request failed?", "Verificar servidor activo, URL/IP correcta, misma red Wi-Fi, puerto 3001 y firewall."),
        ("Importancia de carga/error", "Ayudan a que el usuario sepa si la app esta esperando datos o si ocurrio un problema de conexion."),
    ]
    add_table(doc, ["Pregunta", "Respuesta aplicada al proyecto"], qa_rows, [2750, 6610], header_fill=LIGHT_BLUE, font_size=8.8)

    add_heading(doc, "9. Checklist de cumplimiento", 1)
    checklist = [
        ("API REST propia con JSON Server", "Cumple"),
        ("Recurso con minimo 5 registros y 4 campos", "Cumple: productos con 8 registros base y mas de 4 campos."),
        ("GET mostrado en pantalla movil", "Cumple: catalogo y detalle."),
        ("Carga, error y actualizacion", "Cumple: ActivityIndicator, mensaje de error y RefreshControl."),
        ("POST para registrar un dato", "Cumple: formulario Registrar producto."),
        ("Diseno movil legible", "Cumple: tarjetas, filtros, etiquetas y formulario ordenado."),
        ("Informe con endpoints, evidencias y reflexion", "Cumple: este documento incluye contenido y espacios para capturas."),
    ]
    add_table(doc, ["Criterio", "Estado"], checklist, [4700, 4660], header_fill=LIGHT_GRAY, font_size=8.8)

    add_heading(doc, "10. Publicacion sugerida para foro", 1)
    add_para(
        doc,
        "El principal problema fue el error Network request failed al probar la app en Expo Go. La causa fue que el celular no podia usar localhost y Expo selecciono una IP 169.254.x.x que no correspondia a la red Wi-Fi. La solucion fue levantar JSON Server con --host 0.0.0.0, comprobar desde el navegador del celular que http://192.168.10.110:3001/productos respondia correctamente y configurar EXPO_PUBLIC_API_URL con esa IP. Como mejora, recomiendo siempre probar primero el endpoint en el navegador del telefono antes de depurar el codigo de React Native.",
        size=10.2,
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
